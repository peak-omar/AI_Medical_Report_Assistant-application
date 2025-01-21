import streamlit as st
from dotenv import load_dotenv
from Utils.Agents import Cardiologist, Psychologist, Pulmonologist, MultidisciplinaryTeam
from langchain.vectorstores import FAISS
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.chains import ConversationalRetrievalChain
from langchain.chat_models import ChatOpenAI
from io import BytesIO
from docx import Document
from fpdf import FPDF
import PyPDF2
import os

# Load environment variables
load_dotenv(dotenv_path=".env")

# Streamlit configuration
st.set_page_config(page_title="Medical Report Assistant", layout="wide", page_icon="🏥")

# Initialize session state for documents
if "documents" not in st.session_state:
    st.session_state.documents = {}

if "active_doc" not in st.session_state:
    st.session_state.active_doc = None

# Initialize OpenAI embeddings
embeddings = OpenAIEmbeddings()

# Initialize a function to manage document contexts
def initialize_document_context(doc_name):
    if doc_name not in st.session_state.documents:
        st.session_state.documents[doc_name] = {
            "vector_store": FAISS.from_texts([""], embeddings),
            "conversation_memory": [],
        }

# App header with proper Unicode handling
st.markdown("<h1 style='text-align: center; color: #4CAF50;'>🏥 Medical Report Assistant</h1>", unsafe_allow_html=True)
st.markdown(
    "<p style='text-align: center;'>Upload a patient's medical report, interact with AI specialists, and get a comprehensive diagnosis.</p>",
    unsafe_allow_html=True,
)

# Sidebar file upload
uploaded_file = st.sidebar.file_uploader("📂 Upload Medical Report (TXT, PDF, DOCX)", type=["txt", "pdf", "docx"])

# Functions for file handling
def read_file(file):
    if file.type == "text/plain":
        return file.read().decode("utf-8", errors="replace")
    elif file.type == "application/pdf":
        reader = PyPDF2.PdfReader(file)
        return " ".join([page.extract_text() or "" for page in reader.pages])
    elif file.type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    return None

# Save output functions with UTF-8 handling
def save_output(output_text, file_type):
    if file_type == "TXT":
        return BytesIO(output_text.encode("utf-8", errors="replace"))  # Ensure UTF-8 encoding
    elif file_type == "PDF":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, output_text)
        pdf_file = BytesIO()
        pdf.output(pdf_file)
        pdf_file.seek(0)
        return pdf_file
    elif file_type == "Word":
        doc = Document()
        doc.add_paragraph(output_text)
        word_file = BytesIO()
        doc.save(word_file)
        word_file.seek(0)
        return word_file

# Handle file upload
if uploaded_file:
    medical_report = read_file(uploaded_file)
    if not medical_report:
        st.error("⚠️ Unable to read the file. Please upload a valid TXT, PDF, or DOCX file.")
    else:
        # Assign active document
        doc_name = uploaded_file.name
        st.session_state.active_doc = doc_name

        # Initialize context for the document
        initialize_document_context(doc_name)
        doc_context = st.session_state.documents[doc_name]

        # Display uploaded content
        st.subheader("🔄 Uploaded Medical Report")
        st.text_area("Medical Report Content", medical_report, height=300, disabled=True)

        # Initialize agents
        agents = {
            "Cardiologist": Cardiologist(medical_report),
            "Psychologist": Psychologist(medical_report),
            "Pulmonologist": Pulmonologist(medical_report),
        }

        # Run AI specialists
        if st.button("🩺 Get Specialist Insights"):
            responses = {}
            with st.spinner("Processing medical report..."):
                for role, agent in agents.items():
                    response = agent.run()
                    doc_context["conversation_memory"].append((f"{role} Insights", response))
                    responses[role] = response

                    # Add to vector store
                    doc_context["vector_store"].add_texts([response])

            # Display specialist insights
            st.subheader("🧠 Specialist Insights")
            for role, response in responses.items():
                st.markdown(f"### {role} Response")
                st.write(response)

            # Multidisciplinary team analysis
            st.write("🩻 Running Multidisciplinary Team Analysis...")
            team_agent = MultidisciplinaryTeam(
                cardiologist_report=responses["Cardiologist"],
                psychologist_report=responses["Psychologist"],
                pulmonologist_report=responses["Pulmonologist"],
            )
            final_diagnosis = team_agent.run()
            doc_context["conversation_memory"].append(("Multidisciplinary Team Diagnosis", final_diagnosis))
            doc_context["vector_store"].add_texts([final_diagnosis])

            # Display final diagnosis
            st.subheader("🔍 Multidisciplinary Team Diagnosis")
            st.write(final_diagnosis)

            # Save output options
            st.subheader("💾 Save Diagnosis Report")
            file_type = st.selectbox("Choose File Format", ["TXT", "PDF", "Word"])
            if st.button("📥 Download Report"):
                output_text = f"### Final Diagnosis:\n\n{final_diagnosis}"
                file = save_output(output_text, file_type)
                st.download_button(
                    label=f"Download {file_type} File",
                    data=file,
                    file_name=f"final_diagnosis.{file_type.lower()}",
                    mime="application/octet-stream",
                )

        # Chatbot Interaction
        st.subheader("🤖 Chat with AI Specialists")
        user_query = st.text_input("Ask a question to the AI specialists:")
        if st.button("💬 Ask"):
            if user_query:
                with st.spinner("AI is formulating a response..."):
                    # Retrieve context from FAISS
                    qa_chain = ConversationalRetrievalChain.from_llm(
                        ChatOpenAI(temperature=0.3, model="gpt-4o"),
                        doc_context["vector_store"].as_retriever(),
                    )
                    chat_history = doc_context["conversation_memory"]
                    response = qa_chain.run({"question": user_query, "chat_history": chat_history})
                    doc_context["conversation_memory"].append((user_query, response))
                    st.markdown(f"**User**: {user_query}")
                    st.markdown(f"**AI Specialist**: {response}")

# Document selection for reconnection
if len(st.session_state.documents) > 1:
    st.sidebar.subheader("Reconnect with Previous Documents")
    doc_names = list(st.session_state.documents.keys())
    selected_doc = st.sidebar.selectbox("Select a document to reconnect:", doc_names)
    if st.sidebar.button("Reconnect"):
        st.session_state.active_doc = selected_doc
